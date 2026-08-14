"""Document loading that preserves page boundaries.

Page numbers are the unit of citation in DocuMind, so we keep text grouped
per page (1-based) rather than flattening the whole document. Formats without
physical pages (DOCX/TXT/MD) get pseudo-pages of ~4,000 characters so the
same chunking + citation machinery applies; the UI labels those "part n".
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from pypdf import PdfReader

# Characters per pseudo-page for formats without physical pages (~1K tokens).
_PSEUDO_PAGE_CHARS = 4000

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class PageText:
    page: int          # 1-based
    text: str


def _clean(text: str) -> str:
    """Normalise whitespace and drop hyphenation at line breaks."""
    # join words split across line breaks: "exam-\nple" -> "example"
    text = re.sub(r"-\n(?=\w)", "", text)
    # collapse intra-paragraph single newlines into spaces, keep blank lines
    text = re.sub(r"(?<!\n)\n(?!\n)", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def load_pdf(path: str) -> list[PageText]:
    """Return non-empty pages with cleaned text, in reading order."""
    reader = PdfReader(path)
    pages: list[PageText] = []
    for i, page in enumerate(reader.pages, start=1):
        raw = page.extract_text() or ""
        cleaned = _clean(raw)
        if cleaned:
            pages.append(PageText(page=i, text=cleaned))
    return pages


def _paginate_blocks(blocks: list[str]) -> list[PageText]:
    """Group paragraph blocks into pseudo-pages of ~_PSEUDO_PAGE_CHARS."""
    pages: list[PageText] = []
    cur: list[str] = []
    size = 0
    for block in blocks:
        if size + len(block) > _PSEUDO_PAGE_CHARS and cur:
            pages.append(PageText(page=len(pages) + 1, text="\n\n".join(cur)))
            cur, size = [], 0
        cur.append(block)
        size += len(block)
    if cur:
        pages.append(PageText(page=len(pages) + 1, text="\n\n".join(cur)))
    return pages


def load_docx(path: str) -> list[PageText]:
    """DOCX → pseudo-pages. Heading-styled paragraphs become standalone lines
    so the chunker's section detection fires on them."""
    import docx  # python-docx; imported lazily (only needed for .docx)

    document = docx.Document(path)
    blocks: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "") if para.style else ""
        if style.startswith("Heading") or style == "Title":
            blocks.append(text)  # standalone line → detected as a heading
        else:
            blocks.append(text)
    # Tables: flatten each row to "cell | cell | cell" lines.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return _paginate_blocks(blocks)


def load_text(path: str) -> list[PageText]:
    """TXT/MD → pseudo-pages. Markdown '#' heading markers are stripped to
    plain lines the chunker's heading regex can match."""
    raw = Path(path).read_text(encoding="utf-8", errors="replace")
    if Path(path).suffix.lower() == ".md":
        # "## Setup" -> "Setup"; keep it a standalone line (heading detection).
        raw = re.sub(r"^#{1,6}\s+(.+)$", r"\1", raw, flags=re.MULTILINE)
        # Drop markdown emphasis that would defeat the heading regex.
        raw = re.sub(r"[*_`]{1,3}", "", raw)
    blocks = [b.strip() for b in re.split(r"\n{2,}", raw) if b.strip()]
    return _paginate_blocks(blocks)


def load_document(path: str, filename: str) -> tuple[list[PageText], str]:
    """Dispatch by extension. Returns (pages, source_type)."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return load_pdf(path), "pdf"
    if ext == ".docx":
        return load_docx(path), "docx"
    if ext in (".txt", ".md"):
        return load_text(path), "text"
    raise ValueError(f"Unsupported file type: {ext}")
